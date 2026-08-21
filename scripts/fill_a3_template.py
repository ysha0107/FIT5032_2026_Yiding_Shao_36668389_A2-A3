# -*- coding: utf-8 -*-
"""Fill the A3 Advanced App Report submission template.

Run from D:/Desktop/A2 with the project python:
  PYTHONIOENCODING=utf-8 python mindbridge/scripts/fill_a3_template.py
Output: D:/Desktop/A2/A3-Submission-Yiding-Shao-36668389.docx
"""
import copy
from docx import Document

SRC = 'A3要求/Copy of A3 - Advanced App Report - Submission Template.docx'
OUT = 'A3-Submission-Yiding-Shao-36668389.docx'

TICK = '✅'
# Self-evaluation: rubric columns are Exceeds / Meets / Needs / Fail (1..4)
SELF_EVAL = {
    'D.1': 1, 'D.2': 1, 'D.3': 1, 'D.4': 1,
    'E.1': 1, 'E.2': 1, 'E.3': 2, 'E.4': 1, 'F.1': 1,
}

F1_FEATURES = [
    (
        'Appointment Booking (using Calendar)',
        'FullCalendar week view; clients book weekday 9:00–17:00 slots with '
        'conflict management that blocks double-booked professionals and limits '
        'each user to two upcoming appointments.',
        'Add automated email/SMS reminders and a waiting list, and sync with '
        'staff calendars so cancelled slots reopen automatically.'
    ),
    (
        'Bulk Email',
        'Admin selects multiple users in a searchable, sortable table and sends '
        'each of them a personalised email through the EmailJS service.',
        'Schedule campaigns ahead of time, track open rates, and segment '
        'recipients by role or service history; raise sending quotas.'
    ),
    (
        'Interactive Charts',
        'Chart.js dashboards visualise Firestore data: users by role (doughnut), '
        'average rating per service (bar) and appointments per week (line).',
        'Add time-range filters and drill-down tooltips, plus a client-facing '
        'mood-tracking chart for personalised wellness insights.'
    ),
    (
        'App with at least two offline features',
        'Online/offline banner, Firestore offline persistence, and a localStorage '
        'queue that automatically syncs queued contact messages and ratings when '
        'the connection returns.',
        'Add a service worker for full PWA offline navigation and a conflict-'
        'resolution UI for queued data that clashes during synchronisation.'
    ),
]

HELP_ROWS = [
    ('Claude Code (GenAI tool)',
     'Used for brainstorming the A3 architecture and for debugging during '
     'implementation (explaining errors, suggesting fixes, drafting tests). '
     'Every suggestion was verified against the official documentation and the '
     'running application.'),
    ('Firebase documentation',
     'Authentication flows, Firestore security rules and offline persistence.'),
    ('EmailJS documentation',
     'Sending browser-side emails with file attachments.'),
    ('Leaflet / Nominatim / OSRM documentation',
     'Map rendering, geocoding search and routing APIs.'),
    ('FullCalendar documentation',
     'Calendar views, slot selection and business-hour constraints.'),
    ('Chart.js documentation',
     'Responsive interactive chart configuration.'),
]

REFLECTION = (
    'The most challenging part of this assignment was migrating the A2 app '
    'from localStorage to Firebase while keeping the existing router guards '
    'synchronous — the app now waits for the first Firebase auth-state callback '
    'before mounting, which eliminated a subtle redirect race. Designing global '
    'booking-conflict management across users, and the offline queue that '
    're-syncs with Firestore, also stretched me as a programmer: I had to think '
    'carefully about distributed state, eventual consistency (the Workers KV '
    'list limitation) and graceful degradation. This pushed me well beyond '
    'client-only development into full-stack architecture.'
)


def set_cell(cell, text):
    """Replace cell content with text, keeping the first paragraph's style."""
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r.text = ''
    p.add_run(text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def fill():
    doc = Document(SRC)

    # 1. Declaration paragraph placeholders
    for p in doc.paragraphs:
        if '[Student Name]' in p.text:
            for r in p.runs:
                r.text = r.text.replace('[Student Name]', 'Yiding Shao')
        if '[Assignment Title]' in p.text:
            for r in p.runs:
                r.text = r.text.replace(
                    '[Assignment Title]',
                    'FIT5032 A3 — MindBridge Health Foundation Advanced Web Application (BR D–F)')
        if p.text.strip() == 'Signature: ________________________             		           Date: _____________________________':
            for r in p.runs:
                r.text = 'Signature: Yiding Shao             		           Date: 23 August 2026'

    # 2. GitHub table
    gh = doc.tables[0]
    set_cell(gh.rows[0].cells[1], 'ysha0107')
    set_cell(gh.rows[1].cells[1],
             'https://github.com/ysha0107/FIT5032_2026_Yiding_Shao_36668389_A2\n'
             'Yes — extending the A2 repository (shared with the tutor).')

    # 3. Self-evaluation table
    ev = doc.tables[1]
    for ri in range(1, len(ev.rows)):
        label = ev.rows[ri].cells[0].text
        key = None
        for k in SELF_EVAL:
            if k in label:
                key = k
                break
        if key is None:
            continue
        col = SELF_EVAL[key]
        set_cell(ev.rows[ri].cells[col], TICK)

    # 4. Video link table
    set_cell(doc.tables[2].rows[0].cells[0],
             '<PASTE YOUR GOOGLE DRIVE LINK HERE — share with tutor / Monash accounts>\n'
             'Video: Yiding-Shao-36668389-A3.mp4')

    # 5. F.1 features table (4 rows)
    f1 = doc.tables[3]
    for i, (name, desc, upgrade) in enumerate(F1_FEATURES, start=1):
        cells = f1.rows[i].cells
        set_cell(cells[1], name)
        set_cell(cells[2], desc)
        set_cell(cells[3], upgrade)

    # 6. Reflections: insert after the reflections heading
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith('6. Reflections'):
            # find the next empty paragraph after this heading
            for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                if not doc.paragraphs[j].text.strip():
                    doc.paragraphs[j].text = REFLECTION
                    break
            break

    # 7. Additional help table: fill example row? No — append our rows
    help_tbl = doc.tables[5]
    # fill the empty last row first, then append the rest
    empty_row = help_tbl.rows[-1]
    set_cell(empty_row.cells[0], HELP_ROWS[0][0])
    set_cell(empty_row.cells[1], HELP_ROWS[0][1])
    for name, desc in HELP_ROWS[1:]:
        new_row = copy.deepcopy(help_tbl.rows[-1]._tr)
        help_tbl._tbl.append(new_row)
        row = help_tbl.rows[-1]
        set_cell(row.cells[0], name)
        set_cell(row.cells[1], desc)

    doc.save(OUT)
    print('Saved', OUT)


if __name__ == '__main__':
    fill()
